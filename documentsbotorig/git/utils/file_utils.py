import os
import json
import re
import tempfile
import subprocess
from pathlib import Path
from docxtpl import DocxTemplate
import shutil
from docx import Document as DocxDocument
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from jinja2 import exceptions as jinja2_exceptions
import sys
from copy import deepcopy

# Папка с включёнными шаблонами
ENABLED_PATH = "enabled.json"  # путь к твоему JSON с включёнными шаблонами

BASE_DIR = Path(__file__).resolve().parent.parent
TEMPLATES_DIR = BASE_DIR / "templates"

def load_enabled():
    if not os.path.exists(ENABLED_PATH):
        return {}
    with open(ENABLED_PATH, "r", encoding="utf-8") as f:
        return json.load(f)

def save_enabled(data):
    with open(ENABLED_PATH, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)

def load_templates():
    templates = []
    enabled = load_enabled()
    if not TEMPLATES_DIR.exists():
        raise FileNotFoundError(f"❌ Папка с шаблонами не найдена: {TEMPLATES_DIR}")
    for slug in os.listdir(TEMPLATES_DIR):
        if not enabled.get(slug, False):
            continue
        t_path = TEMPLATES_DIR / slug
        if t_path.is_dir():
            fields_file = t_path / "fields.json"
            if fields_file.exists():
                with open(fields_file, "r", encoding="utf-8") as f:
                    templates.append(json.load(f))
    return templates

def generate_files(template_slug, context, output_dir=None):
    # поддержка нескольких названий шаблонов
    candidates = [
        TEMPLATES_DIR / template_slug / "template.docx",
        TEMPLATES_DIR / template_slug / "template11.docx",
        TEMPLATES_DIR / template_slug / "template1.docx",
    ]
    t_path = next((p for p in candidates if p.exists()), None)
    if not t_path:
        raise FileNotFoundError(
            "❌ Шаблон не найден. Ожидались файлы: "
            + ", ".join(str(p) for p in candidates)
        )

    doc = DocxTemplate(str(t_path))
    try:
        print(f"[tpl] Rendering slug={template_slug}, path={t_path}")
        sys.stdout.flush()
        doc.render(context)
    except jinja2_exceptions.TemplateSyntaxError as e:
        print(f"[tpl] TemplateSyntaxError in slug={template_slug}, path={t_path}: {e}")
        sys.stdout.flush()
        raise

    tmp_docx = tempfile.NamedTemporaryFile(delete=False, suffix=".docx", dir=output_dir)
    doc.save(tmp_docx.name)

    # Пост-обработка: автоматически строим таблицы для полей-массивов
    try:
        fields_cfg_path = TEMPLATES_DIR / template_slug / "fields.json"
        if fields_cfg_path.exists():
            with open(fields_cfg_path, "r", encoding="utf-8") as f:
                cfg = json.load(f)
            arrays = []
            for fld in cfg.get("fields", []):
                if fld.get("type") == "array":
                    items = fld.get("items") or fld.get("item_fields") or []
                    if isinstance(items, list) and items:
                        arrays.append({
                            "key": fld.get("key"),
                            "headers": [i.get("label", i.get("key", "")) for i in items],
                            "col_keys": [i.get("key") for i in items],
                        })
            if arrays:
                _inject_tables_into_docx(tmp_docx.name, arrays, context)
    except Exception as e:
        print(f"⚠ Ошибка автосборки таблиц: {e}")

    pdf_path = None
    soffice_path = shutil.which("soffice")
    if not soffice_path and os.name == "nt":
        possible_path = r"C:\Program Files\LibreOffice\program\soffice.exe"
        if os.path.exists(possible_path):
            soffice_path = possible_path

    if soffice_path:
        try:
            output_dir_final = output_dir or os.path.dirname(tmp_docx.name)
            subprocess.run([
                soffice_path,
                "--headless",
                "--convert-to", "pdf",
                "--outdir", output_dir_final,
                tmp_docx.name
            ], check=True)
            candidate_pdf = os.path.splitext(tmp_docx.name)[0] + ".pdf"
            if os.path.exists(candidate_pdf):
                pdf_path = candidate_pdf
        except subprocess.CalledProcessError as e:
            print(f"⚠ Ошибка при конвертации в PDF: {e}")
    else:
        print("⚠ LibreOffice (soffice) не найден. PDF не будет создан.")

    return tmp_docx.name, pdf_path


def _inject_tables_into_docx(doc_path: str, arrays_meta: list, context: dict):
    """
    Вставляет таблицы для полей-массивов.
    Если в документе есть параграф с текстом {{__TABLE_<key>__}} (или __TABLE_<key>__ / <<TABLE_<key>>>),
    таблица будет вставлена сразу после этого параграфа. Сам параграф не удаляется целиком —
    маркер просто очищается.

    Если маркер не найден — пытаемся найти "якорь" по заголовку (например, содержит "перечень произведений")
    и вставить после него. Если и якоря нет — добавим таблицу в конец документа.
    """
    docx = DocxDocument(doc_path)

    def normalize(text: str):
        return (text or "").strip()

    # Собираем все параграфы верхнего уровня
    paragraphs = list(docx.paragraphs)

    def _try_fill_existing_table(docx_obj: DocxDocument, meta: dict, rows_data: list) -> bool:
        """Пытается заполнить СУЩЕСТВУЮЩУЮ таблицу в шаблоне.
        Возвращает True, если получилось, иначе False (тогда будет создана новая таблица).
        Логика:
        - ищем таблицу, где верхние строки содержат ключевые слова колонок
        - определяем строку начала данных (первая строка, где в первом столбце число '1')
        - клонируем шаблон строки данных и заполняем N строк
        """
        # Более точные ключевые слова с приоритетами
        key_words = {
            "title": {"primary": ["назван"], "secondary": ["произвед", "исполнен", "фонограм"]},
            "music_fio": {"primary": ["автор музык"], "secondary": ["муз", "автор"]},
            "text_fio": {"primary": ["автор текст"], "secondary": ["текст", "слов"]},
            "pseudonym": {"primary": ["исполнител"], "secondary": ["псевдоним", "творческ"]},
            "fonog_fio": {"primary": ["изготовител"], "secondary": ["фонограм", "производител"]},
            "author_rights": {"primary": ["авторск"], "secondary": ["передаваем", "объем"]},
            "neighboring_rights": {"primary": ["смежн"], "secondary": ["передаваем", "объем"]},
            "year": {"primary": ["год"], "secondary": ["выпуск", "дата"]},
        }
        for tbl in docx_obj.tables:
            if not tbl.rows or not tbl.columns:
                continue
            # определяем строку начала данных (первая строка, где первый столбец == '1')
            data_start = None
            for r in range(len(tbl.rows)):
                try:
                    cell0 = (tbl.rows[r].cells[0].text or "").strip()
                except Exception:
                    cell0 = ""
                if cell0 == "1" or cell0.startswith("1"):
                    data_start = r
                    break
            
            # если нашли строку с данными, то заголовки - это всё, что выше неё
            if data_start is not None:
                header_rows_cnt = data_start
            else:
                # если не нашли строку с '1', ищем по заголовкам
                header_rows_cnt = min(4, len(tbl.rows))
                data_start = header_rows_cnt
            
            if data_start >= len(tbl.rows):
                # нет строки данных — добавим одну
                tbl.add_row()
                data_start = len(tbl.rows) - 1
            
            # агрегируем текст только заголовочных строк по колонкам
            cols_n = len(tbl.rows[0].cells)
            col_texts = []
            for c in range(cols_n):
                agg = []
                for r in range(header_rows_cnt):
                    try:
                        agg.append((tbl.rows[r].cells[c].text or "").lower())
                    except Exception:
                        pass
                col_texts.append(" ".join(agg))

            # построить маппинг колонок по заголовкам с приоритетами
            mapping = {}
            print(f"[tpl] Debug: Column texts: {col_texts}")
            
            # Сначала ищем точные соответствия (primary)
            for idx, txt in enumerate(col_texts):
                low = txt.lower()
                for k, kw_dict in key_words.items():
                    if k in mapping or idx in mapping.values():
                        continue
                    primary_kws = kw_dict.get("primary", [])
                    # Проверяем все primary ключевые слова
                    if any(kw in low for kw in primary_kws):
                        mapping[k] = idx
                        print(f"[tpl] Debug: Primary match - mapped column {idx} ('{txt}') to key '{k}'")
                        break
            
            # Затем ищем по secondary, если не нашли primary
            for idx, txt in enumerate(col_texts):
                low = txt.lower()
                for k, kw_dict in key_words.items():
                    if k in mapping or idx in mapping.values():
                        continue
                    secondary_kws = kw_dict.get("secondary", [])
                    # Для secondary требуем больше совпадений или специфичные слова
                    if any(kw in low for kw in secondary_kws):
                        # Дополнительная проверка для конфликтующих полей
                        if k == "text_fio" and "музык" in low:
                            continue
                        if k == "music_fio" and "текст" in low:
                            continue
                        if k in ["pseudonym", "fonog_fio"] and "назван" in low and "title" not in mapping:
                            continue
                        mapping[k] = idx
                        print(f"[tpl] Debug: Secondary match - mapped column {idx} ('{txt}') to key '{k}'")
                        break
            found_cols = len([k for k in meta["col_keys"] if k in mapping])
            if found_cols < 2:
                # слабое совпадение — попробуем использовать стандартный порядок колонок после №
                # 0: №, 1:title, 2:music_fio, 3:text_fio, 4:pseudonym, 5:fonog_fio, 6:author_rights, 7:neighboring_rights, 8:year
                default_order = ["title","music_fio","text_fio","pseudonym","fonog_fio","author_rights","neighboring_rights","year"]
                for i, k in enumerate(default_order, start=1):
                    if i < cols_n:
                        mapping[k] = i
                found_cols = len([k for k in meta["col_keys"] if k in mapping])
                if found_cols < 2:
                    # совсем не похоже на нужную таблицу — ищем дальше
                    continue

            # возьмём шаблон строки данных
            template_row_idx = min(data_start, len(tbl.rows) - 1)
            template_row = tbl.rows[template_row_idx]
            base_tr = deepcopy(template_row._tr)

            # удалим все строки ниже шаблонной (оставим одну шаблонную строку)
            while len(tbl.rows) > template_row_idx + 1:
                tr = tbl.rows[-1]._tr
                tbl._tbl.remove(tr)

            # функция заполнения строки по mapping
            def fill_row(row, i, item):
                cells = row.cells
                if len(cells) > 0:
                    try:
                        cells[0].text = str(i + 1)
                    except Exception as ex:
                        print(f"[tpl] Warning: Could not set row number: {ex}")
                for k in meta["col_keys"]:
                    if k not in mapping:
                        print(f"[tpl] Debug: Key '{k}' not found in mapping")
                        continue
                    col_i = mapping[k]
                    if col_i >= len(cells):
                        print(f"[tpl] Debug: Column index {col_i} out of range (cells: {len(cells)})")
                        continue
                    try:
                        value = str((item or {}).get(k, ""))
                        cells[col_i].text = value
                        print(f"[tpl] Debug: Set cell[{col_i}] (key={k}) = '{value}'")
                    except Exception as ex:
                        print(f"[tpl] Warning: Could not set cell[{col_i}] for key '{k}': {ex}")

            # Проверяем, достаточно ли строк, и нужны ли дополнительные
            if len(rows_data) > 0:
                print(f"[tpl] Debug: Filling {len(rows_data)} rows of data")
                print(f"[tpl] Debug: Template row index: {template_row_idx}")
                print(f"[tpl] Debug: First row data: {rows_data[0]}")
                
                # заполняем первую строку данных (если она есть)
                fill_row(tbl.rows[template_row_idx], 0, rows_data[0])
                
                # добавляем и заполняем остальные строки
                for i in range(1, len(rows_data)):
                    print(f"[tpl] Debug: Adding row {i+1} with data: {rows_data[i]}")
                    # Клонируем шаблонную строку с сохранением форматирования
                    new_tr = deepcopy(base_tr)
                    # Вставляем новую строку после последней строки таблицы
                    tbl._tbl.append(new_tr)
                    # Получаем новую строку и заполняем её
                    new_row = tbl.rows[-1]
                    fill_row(new_row, i, rows_data[i])

            print(f"[tpl] Populated existing table (heuristic) for key={meta['key']}, rows={len(rows_data)}, cols_mapped={found_cols}")
            return True
        return False

    for meta in arrays_meta:
        key = meta["key"]
        rows = context.get(key) or []
        if not isinstance(rows, list) or not rows:
            # нет строк — пропускаем
            continue

        # Сначала пытаемся заполнить существующую таблицу
        try:
            if _try_fill_existing_table(docx, meta, rows):
                print(f"[tpl] Populated existing table for key={key}, rows={len(rows)}")
                continue
        except Exception as e:
            print(f"[tpl] Existing table fill failed for key={key}: {e}")

        # Поддерживаем несколько вариантов маркеров, с нечувствительностью к регистру
        markers = [
            f"{{{{__TABLE_{key}__}}}}",
            f"__TABLE_{key}__",
            f"<<TABLE_{key}>>",
        ]
        markers_low = [m.lower() for m in markers]

        target_para = None
        anchor_para = None
        anchor_keywords = [
            "перечень произведений (таблица)",
            "перечень произведений/исполнений/фонограмм",
            "перечень произведений/исполнений",
            "перечень произведений",
            "перечень работ",
        ]
        anchor_keywords_low = [a.lower() for a in anchor_keywords]

        # Ищем маркер (точное совпадение параграфа)
        for p in paragraphs:
            txt = normalize(p.text)
            low_txt = txt.lower()
            if low_txt in markers_low:
                target_para = p
                break
        # Если не нашли — ищем вхождение токена как подстроки
        if target_para is None:
            for p in paragraphs:
                txt = normalize(p.text)
                low = txt.lower()
                if any(m in low for m in markers_low):
                    target_para = p
                    break
                if anchor_para is None and any(a in low for a in anchor_keywords_low):
                    anchor_para = p

        # Дополнительно ищем якорь внутри таблиц (ячейки)
        if target_para is None and anchor_para is None:
            try:
                for t in docx.tables:
                    for row in t.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                txt = normalize(p.text)
                                low = txt.lower()
                                if any(m in low for m in markers_low):
                                    target_para = p
                                    raise StopIteration
                                if anchor_para is None and any(a in low for a in anchor_keywords_low):
                                    anchor_para = p
                # fallthrough
            except StopIteration:
                pass

        # Определяем точку вставки
        insert_after = target_para or anchor_para

        # Создаём таблицу: +1 колонка для номера
        cols_count = 1 + len(meta["col_keys"])  # No + headers
        table = docx.add_table(rows=1 + len(rows), cols=cols_count)
        try:
            table.style = "Table Grid"
        except Exception:
            pass

        # Заголовок
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "No"
        for i, header in enumerate(meta["headers"], start=1):
            hdr_cells[i].text = str(header)

        # Данные (если нет строк, будет только заголовок)
        for ridx in range(len(rows)):
            row_cells = table.rows[ridx + 1].cells
            row_cells[0].text = str(ridx + 1)
            rdata = rows[ridx]
            for c, k in enumerate(meta["col_keys"], start=1):
                row_cells[c].text = str((rdata or {}).get(k, ""))

        # Перемещение таблицы на место маркера/якоря/в конец
        tbl_el = table._element
        body = docx._element.body
        try:
            body.remove(tbl_el)  # удаляем из конца, чтобы можно было вставить в нужное место
        except Exception:
            pass

        placed = False
        if insert_after is not None:
            try:
                insert_after._element.addnext(tbl_el)
                placed = True
            except Exception as e:
                # Не удалось вставить рядом с параграфом (например, внутри сложной таблицы)
                placed = False
        if not placed:
            # Ни маркера, ни валидного якоря — добавляем в конец документа
            try:
                body.append(tbl_el)
                placed = True
            except Exception:
                placed = False

        # Очистить маркер в параграфе (если он был), без удаления параграфа
        if placed and target_para is not None:
            try:
                txt = target_para.text or ""
                low = txt.lower()
                for m, ml in zip(markers, markers_low):
                    if ml in low:
                        pattern = re.compile(re.escape(m), re.IGNORECASE)
                        txt = pattern.sub("", txt)
                target_para.text = txt
            except Exception:
                pass

        print(f"[tpl] Injected table for key={key}, rows={len(rows)}, placed={'yes' if placed else 'no'}")

    docx.save(doc_path)
