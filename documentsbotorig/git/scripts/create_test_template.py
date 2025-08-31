from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

# Path to target template
BASE = Path(__file__).resolve().parents[1]
TPL_DIR = BASE / "templates" / "add_agreement_OOO"
TPL_DIR.mkdir(parents=True, exist_ok=True)
TPL_PATH = TPL_DIR / "template.docx"


def set_cell_text(cell, text, bold=False, center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    if bold:
        run.bold = True
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font = run.font
    font.name = "Times New Roman"
    r = run._element.rPr.rFonts
    r.set(qn('w:eastAsia'), "Times New Roman")
    font.size = Pt(12)


def build_marker(document):
    # Title
    p = document.add_paragraph()
    run = p.add_run("Перечень произведений/исполнений/фонограмм")
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph()

    # Marker where the bot will inject a table for `works` (safe form, no Jinja braces)
    document.add_paragraph("__TABLE_works__")


def create_template():
    doc = Document()
    # Default paragraph style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    build_marker(doc)

    # Example text below table using single fields
    p = doc.add_paragraph()
    p.add_run("{{ pseudonym }}").bold = True
    p.add_run(" — творческий псевдоним ")
    p.add_run("{{ rodfio }}")

    doc.save(TPL_PATH)
    print(f"Saved: {TPL_PATH}")


if __name__ == "__main__":
    create_template()
